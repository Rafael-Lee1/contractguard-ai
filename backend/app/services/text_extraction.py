import io
from pathlib import Path

import fitz
from docx import Document

from app.utils.text import normalize_whitespace


class ContractTextExtractor:
    SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

    def __init__(self, *, max_upload_size: int):
        self.max_upload_size = max_upload_size

    def validate_upload(self, *, filename: str | None, content: bytes) -> str:
        if not filename:
            raise ValueError("A contract file name is required.")

        extension = Path(filename).suffix.lower()
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError("Unsupported file type. Only PDF and DOCX are accepted.")

        if not content:
            raise ValueError("Uploaded file is empty.")

        if len(content) > self.max_upload_size:
            raise ValueError("Uploaded file exceeds the configured size limit.")

        return extension

    def extract_text(self, *, filename: str, content: bytes) -> str:
        extension = self.validate_upload(filename=filename, content=content)
        if extension == ".pdf":
            return self._extract_from_pdf(content)
        if extension == ".docx":
            return self._extract_from_docx(content)
        raise ValueError("Unsupported file type.")

    def _extract_from_pdf(self, content: bytes) -> str:
        with fitz.open(stream=content, filetype="pdf") as document:
            pages = [page.get_text("text") for page in document]
        text = "\n\n".join(page for page in pages if page.strip())
        normalized = normalize_whitespace(text)
        if not normalized:
            raise ValueError("No readable text was found in the PDF.")
        return normalized

    def _extract_from_docx(self, content: bytes) -> str:
        document = Document(io.BytesIO(content))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

        table_cells: list[str] = []
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        table_cells.append(cell_text)

        text = "\n\n".join(paragraphs + table_cells)
        normalized = normalize_whitespace(text)
        if not normalized:
            raise ValueError("No readable text was found in the DOCX file.")
        return normalized
